#!/usr/bin/env python3
"""
MDX Converter - Convert MDX files to PDF and/or DOCX formats.
"""

import argparse
import os
import sys
import re
from pathlib import Path

def convert_mdx_to_html(mdx_path):
    """Convert MDX content to HTML."""
    try:
        with open(mdx_path, 'r', encoding='utf-8') as f:
            mdx_content = f.read()
        
        # Strip out JSX/React components from MDX (simplified approach)
        mdx_content = re.sub(r'import.*?;', '', mdx_content)
        mdx_content = re.sub(r'<[A-Z][^>]*>.*?</[A-Z][^>]*>', '', mdx_content, flags=re.DOTALL)
        mdx_content = re.sub(r'<[A-Z][^/>]*?/>', '', mdx_content)
        
        # Convert Markdown to HTML
        try:
            from markdown_it import MarkdownIt
        except ImportError:
            print("Error: markdown-it-py is required. Install it using 'pip install markdown-it-py'", file=sys.stderr)
            sys.exit(1)
            
        md = MarkdownIt()
        html_content = md.render(mdx_content)
        
        return html_content
    except Exception as e:
        print(f"Error converting MDX to HTML: {e}", file=sys.stderr)
        sys.exit(1)

def convert_mdx_to_pdf(mdx_path, output_path):
    """Convert MDX file to PDF."""
    try:
        try:
            import markdown2
            from xhtml2pdf import pisa
        except ImportError:
            print("Error: markdown2 and xhtml2pdf are required for PDF conversion.", file=sys.stderr)
            print("Install them using 'pip install markdown2 xhtml2pdf'", file=sys.stderr)
            sys.exit(1)
            
        # Read and clean MDX content
        with open(mdx_path, 'r', encoding='utf-8') as f:
            mdx_content = f.read()
        
        # Strip out JSX/React components from MDX
        mdx_content = re.sub(r'import.*?;', '', mdx_content)
        mdx_content = re.sub(r'<[A-Z][^>]*>.*?</[A-Z][^>]*>', '', mdx_content, flags=re.DOTALL)
        mdx_content = re.sub(r'<[A-Z][^/>]*?/>', '', mdx_content)
        
        # Convert to HTML
        html_content = markdown2.markdown(mdx_content)
        
        # Add basic HTML structure
        html_document = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <style>
                body {{ font-family: Arial, sans-serif; margin: 20px; }}
                pre {{ background-color: #f5f5f5; padding: 10px; border-radius: 5px; }}
                code {{ font-family: monospace; }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """
        
        # Convert HTML to PDF
        with open(output_path, "wb") as pdf_file:
            pisa_status = pisa.CreatePDF(html_document, dest=pdf_file)
            
        if pisa_status.err:
            print(f"Error converting HTML to PDF: {pisa_status.err}", file=sys.stderr)
            sys.exit(1)
            
        print(f"Created PDF: {output_path}")
    except Exception as e:
        print(f"Error converting to PDF: {e}", file=sys.stderr)
        sys.exit(1)

def convert_mdx_to_docx(mdx_path, output_path):
    """Convert MDX file to DOCX."""
    try:
        try:
            import markdown2
            from docx import Document
            from docx.shared import Pt
            from bs4 import BeautifulSoup
        except ImportError:
            print("Error: markdown2, python-docx, and beautifulsoup4 are required for DOCX conversion.", file=sys.stderr)
            print("Install them using 'pip install markdown2 python-docx beautifulsoup4'", file=sys.stderr)
            sys.exit(1)
            
        # Read and clean MDX content
        with open(mdx_path, 'r', encoding='utf-8') as f:
            mdx_content = f.read()
        
        # Strip out JSX/React components from MDX
        mdx_content = re.sub(r'import.*?;', '', mdx_content)
        mdx_content = re.sub(r'<[A-Z][^>]*>.*?</[A-Z][^>]*>', '', mdx_content, flags=re.DOTALL)
        mdx_content = re.sub(r'<[A-Z][^/>]*?/>', '', mdx_content)
        
        # Convert to HTML
        html_content = markdown2.markdown(mdx_content, extras=['fenced-code-blocks', 'tables'])
        
        # Parse HTML
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Create a new Word document
        doc = Document()
        
        # Process HTML elements and add to document
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'pre', 'code', 'ul', 'ol', 'li']):
            if element.name.startswith('h'):
                level = int(element.name[1])
                doc.add_heading(element.text, level=level)
            elif element.name == 'p':
                doc.add_paragraph(element.text)
            elif element.name == 'pre' or element.name == 'code':
                # Add code blocks with a different style
                p = doc.add_paragraph(element.text)
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
            elif element.name == 'ul':
                for li in element.find_all('li', recursive=False):
                    doc.add_paragraph('• ' + li.text, style='List Bullet')
            elif element.name == 'ol':
                for i, li in enumerate(element.find_all('li', recursive=False), 1):
                    doc.add_paragraph(f"{i}. {li.text}", style='List Number')
        
        # Save the document
        doc.save(output_path)
        print(f"Created DOCX: {output_path}")
    except Exception as e:
        print(f"Error converting to DOCX: {e}", file=sys.stderr)
        sys.exit(1)

def process_file(mdx_path, output_dir, formats):
    """Process a single MDX file according to the requested formats."""
    filename = mdx_path.stem  # Get filename without extension
    
    for fmt in formats:
        output_path = os.path.join(output_dir, f"{filename}.{fmt}")
        
        if fmt == 'pdf':
            convert_mdx_to_pdf(mdx_path, output_path)
        elif fmt == 'docx':
            convert_mdx_to_docx(mdx_path, output_path)

def process_directory(mdx_dir, output_dir, formats):
    """Process all MDX files in a directory."""
    mdx_files = list(Path(mdx_dir).glob('**/*.mdx'))
    
    if not mdx_files:
        print(f"No MDX files found in '{mdx_dir}'")
        return
    
    print(f"Found {len(mdx_files)} MDX files to process")
    for mdx_file in mdx_files:
        process_file(mdx_file, output_dir, formats)

def main():
    """Main entry point for the script."""
    parser = argparse.ArgumentParser(description='Convert MDX files to PDF and/or DOCX formats.')
    parser.add_argument('input_path', help='Path to MDX file or directory containing MDX files')
    parser.add_argument('output_directory', help='Path to output directory')
    parser.add_argument('--format', choices=['pdf', 'docx', 'both'], default='both',
                       help='Output format (default: both)')
    
    args = parser.parse_args()
    
    # Convert format to a list of formats
    formats = []
    if args.format == 'both':
        formats = ['pdf', 'docx']
    else:
        formats = [args.format]
    
    # Create output directory if it doesn't exist
    os.makedirs(args.output_directory, exist_ok=True)
    
    # Determine if input is a file or directory
    input_path = Path(args.input_path)
    if input_path.is_file():
        if input_path.suffix.lower() != '.mdx':
            print(f"Error: '{input_path}' is not an MDX file.", file=sys.stderr)
            sys.exit(1)
        process_file(input_path, args.output_directory, formats)
    elif input_path.is_dir():
        process_directory(input_path, args.output_directory, formats)
    else:
        print(f"Error: '{input_path}' not found or is not a file/directory.", file=sys.stderr)
        sys.exit(1)

if __name__ == '__main__':
    main()